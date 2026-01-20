"""
Resume text extraction from PDF, DOCX, and MD formats.

Usage:
    from backend.extractors import extract_text
    
    text = extract_text("resume.pdf")  # Works with .pdf, .docx, .md
"""

from pathlib import Path
from typing import Union

import fitz  # PyMuPDF
from docx import Document


def extract_from_pdf(file_path: Union[str, Path]) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Uses block-based extraction to preserve layout structure.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    text_blocks = []
    
    with fitz.open(file_path) as doc:
        for page in doc:
            # Get text blocks with position info
            blocks = page.get_text("blocks")
            
            # Sort by vertical position, then horizontal
            blocks.sort(key=lambda b: (b[1], b[0]))
            
            for block in blocks:
                # block[4] is the text content
                if block[6] == 0:  # Text block (not image)
                    text_blocks.append(block[4].strip())
    
    return "\n".join(text_blocks)


def extract_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Preserves paragraph structure.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    doc = Document(file_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))
    
    return "\n".join(paragraphs)


def extract_from_md(file_path: Union[str, Path]) -> str:
    """
    Read text from a Markdown file.
    
    Args:
        file_path: Path to the MD file
        
    Returns:
        File content as string
    """
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_text(file_path: Union[str, Path]) -> str:
    """
    Extract text from a resume file (PDF, DOCX, or MD).
    
    Automatically detects format based on file extension.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file does not exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    suffix = path.suffix.lower()
    
    extractors = {
        ".pdf": extract_from_pdf,
        ".docx": extract_from_docx,
        ".md": extract_from_md,
        ".markdown": extract_from_md,
    }
    
    if suffix not in extractors:
        supported = ", ".join(extractors.keys())
        raise ValueError(f"Unsupported format: {suffix}. Supported: {supported}")
    
    return extractors[suffix](path)


def parse_resume_preview(text: str) -> dict:
    """
    Quick preview of detected sections without full parsing.
    
    Returns summary of what was detected.
    """
    lines = text.strip().split("\n")
    
    # Simple heuristics for preview
    preview = {
        "total_lines": len(lines),
        "detected_sections": [],
        "sample_content": text[:500] + "..." if len(text) > 500 else text,
    }
    
    # Detect common section headers
    section_keywords = ["work", "experience", "education", "skill", "project", "summary"]
    for line in lines:
        lower = line.lower().strip()
        for keyword in section_keywords:
            if keyword in lower and len(lower) < 30:
                if keyword not in preview["detected_sections"]:
                    preview["detected_sections"].append(keyword)
    
    return preview


if __name__ == "__main__":
    # Quick test
    import sys
    
    if len(sys.argv) > 1:
        text = extract_text(sys.argv[1])
        print(f"Extracted {len(text)} characters:")
        print("-" * 40)
        print(text[:500])
    else:
        print("Usage: python extractors.py <file_path>")
